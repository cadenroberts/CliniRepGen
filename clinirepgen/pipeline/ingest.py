"""
Ingest Stage - Builds Trial Manifest from source documents.

Handles:
- Document discovery and loading
- PDF/Word/text parsing
- Table extraction
- ClinicalTrials.gov data integration
"""

import os
import json
import logging
from typing import List, Optional, Dict, Any
from pathlib import Path
from datetime import datetime

from clinirepgen.manifest.models import TrialManifest, DocumentType
from clinirepgen.manifest.builder import ManifestBuilder

logger = logging.getLogger(__name__)


class IngestStage:
    """
    Handles document ingestion into Trial Manifest.
    
    Supports:
    - PDF documents (protocols, CSRs, papers)
    - Text files
    - JSON data (CT.gov exports)
    - Word documents
    """
    
    def __init__(self, trial_id: str, output_dir: str = "."):
        """
        Initialize the ingest stage.
        
        Args:
            trial_id: Unique trial identifier
            output_dir: Directory for output files
        """
        self.trial_id = trial_id
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        self.builder = ManifestBuilder(
            trial_id=trial_id,
            output_dir=str(output_dir),
        )
    
    def ingest_directory(
        self,
        directory: str,
        file_types: Optional[List[str]] = None,
    ) -> TrialManifest:
        """
        Ingest all supported documents from a directory.
        
        Args:
            directory: Path to directory containing documents
            file_types: Optional list of extensions to include (e.g., ['.pdf', '.txt'])
            
        Returns:
            Built TrialManifest
        """
        dir_path = Path(directory)
        if not dir_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        default_types = ['.pdf', '.txt', '.json', '.docx', '.doc', '.md']
        allowed_types = file_types or default_types
        
        # Find all matching files
        files = []
        for ext in allowed_types:
            files.extend(dir_path.glob(f"*{ext}"))
            files.extend(dir_path.glob(f"**/*{ext}"))  # Recursive
        
        logger.info(f"Found {len(files)} documents to ingest")
        
        # Ingest each file, validating it is inside the requested directory
        dir_resolved = str(dir_path.resolve())
        for file_path in sorted(set(files)):
            try:
                fp = Path(file_path).resolve()
                # Prevent path traversal / symlink escape: ensure file is inside directory
                if not str(fp).startswith(dir_resolved):
                    logger.warning(f"Skipping file outside directory: {file_path}")
                    continue
                self.ingest_file(str(fp), base_dir=dir_resolved)
            except Exception as e:
                logger.error(f"Failed to ingest {file_path}: {e}")
        
        return self.builder.build()
    
    def ingest_file(self, file_path: str, base_dir: Optional[str] = None) -> str:
        """
        Ingest a single file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Document ID
        """
        path = Path(file_path)
        # If a base_dir is provided, validate that the file is contained within it
        if base_dir is not None:
            try:
                if not str(path.resolve()).startswith(str(Path(base_dir).resolve())):
                    raise PermissionError(f"File {file_path} is outside of allowed base directory")
            except Exception:
                raise PermissionError(f"Invalid or disallowed file path: {file_path}")
        ext = path.suffix.lower()
        
        # Add document to manifest
        doc_id = self.builder.add_document(file_path)
        
        # Parse based on file type
        if ext == '.txt' or ext == '.md':
            self._parse_text_file(doc_id, file_path)
        elif ext == '.json':
            self._parse_json_file(doc_id, file_path)
        elif ext == '.pdf':
            self._parse_pdf_file(doc_id, file_path)
        elif ext in ['.docx', '.doc']:
            self._parse_word_file(doc_id, file_path)
        else:
            logger.warning(f"Unsupported file type: {ext}")
        
        return doc_id
    
    def ingest_ctgov(self, ctgov_json: Dict[str, Any]) -> str:
        """
        Ingest ClinicalTrials.gov data.
        
        Args:
            ctgov_json: Parsed CT.gov JSON data
            
        Returns:
            Document ID
        """
        return self.builder.add_ctgov_data(ctgov_json)
    
    def ingest_ctgov_file(self, file_path: str) -> str:
        """
        Ingest CT.gov data from a JSON file.
        
        Args:
            file_path: Path to JSON file
            
        Returns:
            Document ID
        """
        with open(file_path, "r") as f:
            data = json.load(f)
        return self.ingest_ctgov(data)
    
    def _parse_text_file(self, doc_id: str, file_path: str) -> None:
        """Parse a text file into sections."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        
        self.builder.process_document_text(doc_id, text)
    
    def _parse_json_file(self, doc_id: str, file_path: str) -> None:
        """Parse a JSON file."""
        with open(file_path, "r") as f:
            data = json.load(f)
        
        # Check if it's CT.gov format
        if "study" in data or "nct_id" in data:
            # Re-ingest as CT.gov data
            self.builder.add_ctgov_data(data)
        else:
            # Treat as general JSON - convert to text
            text = json.dumps(data, indent=2)
            self.builder.process_document_text(doc_id, text)
    
    def _parse_pdf_file(self, doc_id: str, file_path: str) -> None:
        """Parse a PDF file using pdfplumber."""
        try:
            import pdfplumber
            
            full_text = []
            table_index = 0
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    # Extract text
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table and len(table) > 1:
                            self._add_table(doc_id, table, page_num, table_index)
                            table_index += 1
            
            # Process full text
            if full_text:
                self.builder.process_document_text(doc_id, "\n\n".join(full_text))
            
            # Update document metadata
            doc = self.builder.documents.get(doc_id)
            if doc:
                doc.num_pages = len(pdf.pages) if 'pdf' in dir() else None
                
        except ImportError:
            logger.warning("pdfplumber not available, skipping PDF parsing")
        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}")
    
    def _parse_word_file(self, doc_id: str, file_path: str) -> None:
        """Parse a Word document using python-docx."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables
            table_index = 0
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    rows.append(cells)
                
                if rows:
                    self._add_table(doc_id, rows, 1, table_index)
                    table_index += 1
            
            # Process text
            if paragraphs:
                self.builder.process_document_text(doc_id, "\n\n".join(paragraphs))
                
        except ImportError:
            logger.warning("python-docx not available, skipping Word parsing")
        except Exception as e:
            logger.error(f"Failed to parse Word doc {file_path}: {e}")
    
    def _add_table(
        self,
        doc_id: str,
        table_data: List[List[str]],
        page_num: int,
        table_index: int,
    ) -> str:
        """Add an extracted table to the manifest."""
        
        # First row is typically headers
        headers = table_data[0] if table_data else []
        
        table_dict = {
            "raw_data": table_data,
            "headers": headers,
            "page_num": page_num,
            "caption": f"Table {table_index + 1}",
            "table_number": str(table_index + 1),
        }
        
        return self.builder.add_table(doc_id, table_dict)
    
    def build(self) -> TrialManifest:
        """Build and return the final manifest."""
        return self.builder.build()
    
    def save(self, output_path: Optional[str] = None) -> str:
        """Save the manifest to a JSON file."""
        return self.builder.save(output_path)
